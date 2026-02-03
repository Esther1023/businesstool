from docx import Document
from docxtpl import DocxTemplate
from typing import Dict, List, Optional
import re
import os
import tempfile
from datetime import datetime, timedelta

def convert_to_lowercase(text: str) -> str:
    """
    将中文大写数字转换为小写
    """
    mapping = {
        '零': '〇', '壹': '一', '贰': '二', '叁': '三', '肆': '四',
        '伍': '五', '陆': '六', '柒': '七', '捌': '八', '玖': '九',
        '拾': '十', '佰': '百', '仟': '千', '万': '万', '亿': '亿',
        '圆': '元'
    }
    result = text
    for upper, lower in mapping.items():
        result = result.replace(upper, lower)
    return result

class TemplateHandler:
    def __init__(self, template_path: str):
        """
        初始化模板处理器
        
        参数:
            template_path (str): Word模板文件路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        self.template_path = template_path
        try:
            self.doc = DocxTemplate(template_path)
        except Exception as e:
            # 尝试复制文件到临时目录再加载
            temp_path = os.path.join(tempfile.gettempdir(), os.path.basename(template_path))
            with open(template_path, 'rb') as src, open(temp_path, 'wb') as dst:
                dst.write(src.read())
            self.doc = DocxTemplate(temp_path)
            os.unlink(temp_path)  # 清理临时文件

    def process_template(self, context: Dict[str, str], output_path: Optional[str] = None, lowercase_amount: bool = True) -> str:
        """
        处理模板并保存为新文件
        
        参数:
            context (Dict[str, str]): 要替换的变量字典
            output_path (str, optional): 新文件保存路径，如果为None则自动生成
            
        返回:
            str: 保存的文件路径
        """
        # 检查是否为报价单模板
        is_quote_template = '报价单' in self.template_path
        
        if is_quote_template:
            # 报价单专用处理逻辑 - 只使用6个核心字段
            context = self._process_quote_template(context)
        else:
            # 合同模板处理逻辑
            context = self._process_contract_template(context)
        
        # 渲染模板
        self.doc.render(context)

        # 生成输出路径
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.splitext(os.path.basename(self.template_path))[0]
            output_path = f"{filename}_{timestamp}.docx"

        # 先保存到临时路径
        tmp_path = os.path.join(tempfile.gettempdir(), f"_tmp_{os.path.basename(output_path)}")
        self.doc.save(tmp_path)

        # 尝试回填付款日期到含有“年 月 日”占位的段落/单元格
        try:
            py_doc = Document(tmp_path)
            pay = f"{context.get('payment_year','') }年{context.get('payment_month','') }月{context.get('payment_day','') }日"
            if pay.replace(' ', '') != '年月日':
                self._inject_payment_date_fallback(py_doc, pay)
            py_doc.save(output_path)
            # 移除XML级别的全局兜底，避免误伤服务期限
            # self._inject_payment_date_xml(output_path, pay)
        except Exception:
            # 回退为直接保存渲染结果
            self.doc.save(output_path)
        finally:
            if os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

        return output_path
    
    def _process_quote_template(self, context: Dict[str, str]) -> Dict[str, str]:
        """
        处理报价单模板的变量
        只使用6个核心字段：company_name, tax_number, jdy_account, total_amount, user_count, unit_price
        """
        # 核心字段保持不变
        processed_context = {
            'company_name': context.get('company_name', ''),
            'tax_number': context.get('tax_number', ''),
            'jdy_account': context.get('jdy_account', ''),
            'total_amount': context.get('total_amount', '0'),
            'user_count': context.get('user_count', '1'),
            'unit_price': context.get('unit_price', '0'),
        }
        
        # 为其他可能的模板变量提供默认值
        default_values = {
            # 基本信息默认值
            'reg_address': '详见营业执照',
            'reg_phone': '详见营业执照', 
            'bank_name': '详见开户许可证',
            'bank_account': '详见开户许可证',
            'contact_name': '详见合同',
            'contact_phone': '详见合同',
            'mail_address': '详见合同',
            
            # 服务期限默认值（当前年度）
            'service_years': '1',
            'start_year': str(datetime.now().year),
            'start_month': str(datetime.now().month),
            'start_day': str(datetime.now().day),
            'end_year': str(datetime.now().year + 1),
            'end_month': str(datetime.now().month),
            'end_day': str(datetime.now().day),
            
            # 费用信息
            'total_amount_cn': self._number_to_chinese(processed_context['total_amount']),
            'payment_amount': processed_context['total_amount'],
            'payment_amount_cn': self._number_to_chinese(processed_context['total_amount']),
            'tax_rate': '6%',
            
            # 发票信息
            'invoice_type': '普通发票',
            
            # 合同相关
            'contract_types': '续费服务',
            'second_row': 'false',
            'table_row_count': '1'
        }
        
        # 合并默认值
        processed_context.update(default_values)

        # 计算到期日期减3天并拆分为年/月/日
        try:
            ey = int(str(context.get('end_year', processed_context['end_year'])))
            em = int(str(context.get('end_month', processed_context['end_month'])))
            ed = int(str(context.get('end_day', processed_context['end_day'])))
            end_dt = datetime(ey, em, ed)
            payment_dt = end_dt - timedelta(days=3)
            processed_context['payment_year'] = str(payment_dt.year)
            processed_context['payment_month'] = str(payment_dt.month)
            processed_context['payment_day'] = str(payment_dt.day)
        except Exception:
            # 回退到当天-3天
            fallback_dt = datetime.now() - timedelta(days=3)
            processed_context['payment_year'] = str(fallback_dt.year)
            processed_context['payment_month'] = str(fallback_dt.month)
            processed_context['payment_day'] = str(fallback_dt.day)

        # today（YYYY年MM月DD日）
        try:
            processed_context['today'] = datetime.now().strftime('%Y年%m月%d日')
        except Exception:
            processed_context['today'] = ''
        
        return processed_context
    
    def _process_contract_template(self, context: Dict[str, str]) -> Dict[str, str]:
        """
        处理合同模板的变量（保持原有逻辑）
        """
        # 设置固定税率（带百分号）
        context['tax_rate'] = '6%'
        
        # 处理合同类型
        contract_types = context.get('contract_types', ['续费'])
        if isinstance(contract_types, str):
            contract_types = [contract_types]
        context['contract_types'] = ', '.join(contract_types)
        
        # 处理表格行数
        table_row_count = int(context.get('table_row_count', '1'))
        if table_row_count == 2:
            # 复制第一行的数据到第二行
            context['second_row'] = 'true'
        else:
            context['second_row'] = 'false'
            
        # 付款日期（到期日前3天）与 today
        try:
            ey = int(str(context.get('end_year', datetime.now().year)))
            em = int(str(context.get('end_month', datetime.now().month)))
            ed = int(str(context.get('end_day', datetime.now().day)))
            end_dt = datetime(ey, em, ed)
            payment_dt = end_dt - timedelta(days=3)
            context['payment_year'] = str(payment_dt.year)
            context['payment_month'] = str(payment_dt.month)
            context['payment_day'] = str(payment_dt.day)
        except Exception:
            fallback_dt = datetime.now() - timedelta(days=3)
            context['payment_year'] = str(fallback_dt.year)
            context['payment_month'] = str(fallback_dt.month)
            context['payment_day'] = str(fallback_dt.day)

        try:
            context['today'] = datetime.now().strftime('%Y年%m月%d日')
        except Exception:
            context['today'] = ''

        # 金额字段与中文金额的兜底
        if not context.get('payment_amount'):
            context['payment_amount'] = context.get('total_amount', '0')
        if not context.get('payment_amount_cn'):
            context['payment_amount_cn'] = self._number_to_chinese(context.get('payment_amount', '0'))
        # 模板别名容错（空格写法）
        context['payment amount cn'] = context.get('payment_amount_cn')

        return context

    def _inject_payment_date_fallback(self, py_doc: Document, date_str: str) -> None:
        """
        在渲染后文档中，将形如“年 _ 月 __ 日”或“年  月   日”的占位，替换为指定日期字符串。
        仅在包含“付款方式”或“甲方应在”的段落中执行，避免误伤服务期限。
        """
        import re

        patterns = [
            re.compile(r"年\s*_+\s*月\s*_+\s*日"),
            re.compile(r"年\s+月\s+日"),
            re.compile(r"\{\{\s*payment_year\s*\}\}\s*年\s*\{\{\s*payment_month\s*\}\}\s*月\s*\{\{\s*payment_day\s*\}\}\s*日"),
        ]

        def replace_in_paragraph(p):
            # 仅处理包含“付款方式”或“甲方应在”的段落，避免误伤服务期限
            if '付款方式' not in p.text and '甲方应在' not in p.text:
                return False

            text = p.text
            for pat in patterns:
                if pat.search(text):
                    new_text = pat.sub(date_str, text, count=1)
                    # 只有在真正发生替换时才更新 paragraph.text，避免不必要的格式丢失
                    if new_text != text:
                        p.text = new_text
                        return True
            return False

        # 段落
        for p in py_doc.paragraphs:
            replace_in_paragraph(p)

        # 表格单元格
        for table in py_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

    def _inject_payment_date_xml(self, docx_path: str, date_str: str) -> None:
        """
        使用正则在 XML 层进行兜底替换，覆盖被 run/tag 分割的“年…月…日”序列。
        """
        import zipfile, tempfile, shutil, re

        try:
            with zipfile.ZipFile(docx_path, 'r') as zin:
                xml = zin.read('word/document.xml').decode('utf-8', 'ignore')

            # 匹配被标签/样式分割的日期占位：年…月…日（尽量局部、非贪婪）
            patterns = [
                r"年(?:\s|_|\u3000|</w:t>.*?<w:t>){0,20}月(?:\s|_|\u3000|</w:t>.*?<w:t>){0,20}日",
                r"\{\{\s*payment_year\s*\}\}(?:</w:t>.*?<w:t>|\s){0,20}年(?:</w:t>.*?<w:t>|\s){0,20}\{\{\s*payment_month\s*\}\}(?:</w:t>.*?<w:t>|\s){0,20}月(?:</w:t>.*?<w:t>|\s){0,20}\{\{\s*payment_day\s*\}\}(?:</w:t>.*?<w:t>|\s){0,20}日",
            ]

            new_xml = xml
            replaced = False
            for pat in patterns:
                new_xml2 = re.sub(pat, date_str, new_xml, count=1, flags=re.DOTALL)
                if new_xml2 != new_xml:
                    new_xml = new_xml2
                    replaced = True
                    break

            if not replaced:
                return

            tmp_zip = tempfile.mktemp(suffix='.docx')
            with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(tmp_zip, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == 'word/document.xml':
                        zout.writestr(item, new_xml)
                    else:
                        zout.writestr(item, zin.read(item.filename))

            shutil.move(tmp_zip, docx_path)
        except Exception:
            pass
    
    def _number_to_chinese(self, amount_str: str) -> str:
        """
        将数字金额转换为中文大写
        """
        try:
            amount = float(amount_str)
            if amount == 0:
                return "零圆整"
            
            # 简化的中文数字转换
            units = ['', '十', '百', '千', '万']
            digits = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']
            
            # 处理整数部分
            integer_part = int(amount)
            if integer_part == 0:
                return "零圆整"
            
            # 简单转换逻辑（适用于常见金额）
            if integer_part < 10:
                return f"{digits[integer_part]}圆整"
            elif integer_part < 100:
                tens = integer_part // 10
                ones = integer_part % 10
                if ones == 0:
                    return f"{digits[tens]}十圆整"
                else:
                    return f"{digits[tens]}十{digits[ones]}圆整"
            else:
                # 对于更大的数字，返回简化格式
                return f"{integer_part}圆整"
                
        except (ValueError, TypeError):
            return "零圆整"

    def get_template_variables(self) -> List[str]:
        """
        获取模板中的所有变量
        
        返回:
            List[str]: 变量名列表
        """
        # 定义合同中所有需要替换的变量
        variables = {
            # 甲方基本信息（必填项标记）
            'company_name': '公司名称（甲方）（必填）',
            'tax_number': '税号（必填）',
            'reg_address': '注册地址（必填）',
            'reg_phone': '注册电话（必填）',
            'bank_name': '开户行',
            'bank_account': '账号',
            'contact_name': '联系人',
            'contact_phone': '联系电话',
            'mail_address': '邮寄地址',
            'jdy_account': '简道云账号（必填）',
            
            # 服务期限（全部必填）
            'service_years': '服务年限（必填）',
            'start_year': '开始年份（必填）',
            'start_month': '开始月份（必填）',
            'start_day': '开始日期（必填）',
            'end_year': '结束年份（自动计算）',
            'end_month': '结束月份（自动计算）',
            'end_day': '结束日期（自动计算）',
            
            # 费用信息
            'unit_price': '单价（元/人/年）',
            'total_amount': '服务费用金额（数字）',
            'total_amount_cn': '服务费用金额（大写，自动生成）',
            'payment_amount': '服务费用金额（数字，同上）',
            'payment_amount_cn': '服务费用金额（大写，自动生成）',
            'tax_rate': '税率（固定为6%）',
            
            # 用户信息
            'user_count': '使用人数',
            
            # 发票信息
            'invoice_type': '发票类型（普通/专用）'
        }
        
        return sorted(list(variables))
